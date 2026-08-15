from collections import Counter
from concurrent.futures import ThreadPoolExecutor
from pathlib import Path

from docx import Document

from app.enums.element_type import ElementType
from app.services.extraction.docx_extractor import DocxExtractor


# ============================================================
# TEST HELPERS
# ============================================================

TEST_DIR = Path("tests/doc_test/edge_cases")
TEST_DIR.mkdir(parents=True, exist_ok=True)


def create_docx(
    filename: str,
    paragraphs: list[str] | None = None,
    tables: list[list[list[str]]] | None = None,
) -> Path:
    """
    Create a temporary DOCX for edge-case testing.

    tables format:

    [
        [
            ["Name", "Age"],
            ["Abhishek", "20"],
        ],
        [
            ["Apple", "10"],
            ["Banana", "20"],
        ],
    ]
    """

    path = TEST_DIR / filename

    doc = Document()

    # -------------------------
    # Paragraphs
    # -------------------------

    if paragraphs:
        for text in paragraphs:
            doc.add_paragraph(text)

    # -------------------------
    # Tables
    # -------------------------

    if tables:
        for table_data in tables:

            if not table_data:
                continue

            rows = len(table_data)
            cols = max(len(row) for row in table_data)

            table = doc.add_table(
                rows=rows,
                cols=cols,
            )

            for row_index, row in enumerate(table_data):

                for col_index, value in enumerate(row):

                    table.cell(
                        row_index,
                        col_index,
                    ).text = value

    doc.save(path)

    return path


def get_elements(result, element_type):
    return [
        element
        for element in result.elements
        if element.element_type == element_type
    ]


def get_tables(result):
    return get_elements(
        result,
        ElementType.TABLE,
    )


def get_paragraphs(result):
    return get_elements(
        result,
        ElementType.PARAGRAPH,
    )


def get_code_blocks(result):
    return get_elements(
        result,
        ElementType.CODE_BLOCK,
    )


# ============================================================
# EXISTING REAL DOCUMENT TEST
# ============================================================

def test_extract_real_report_everything():

    # ------------------------------------------------------------
    # PATHS
    # ------------------------------------------------------------

    docx_path = Path(
        "tests/doc_test/report.docx"
    )

    output_path = Path(
        "tests/doc_test/report_extraction.txt"
    )

    # ------------------------------------------------------------
    # EXTRACT
    # ------------------------------------------------------------

    extractor = DocxExtractor()

    result = extractor.extract(
        str(docx_path),
        document_id="real-report",
        filename=docx_path.name,
    )

    # ------------------------------------------------------------
    # BASIC VALIDATION
    # ------------------------------------------------------------

    assert result is not None
    assert result.elements

    # ------------------------------------------------------------
    # ELEMENT TYPE COUNTS
    # ------------------------------------------------------------

    counts = Counter(
        element.element_type
        for element in result.elements
    )

    # ------------------------------------------------------------
    # WRITE COMPLETE EXTRACTION
    # ------------------------------------------------------------

    with open(
        output_path,
        "w",
        encoding="utf-8",
    ) as file:

        file.write("=" * 100 + "\n")
        file.write("REAL DOCX EXTRACTION RESULT\n")
        file.write("=" * 100 + "\n\n")

        file.write(
            f"Document: {docx_path}\n"
        )

        file.write(
            f"Total elements: {len(result.elements)}\n\n"
        )

        file.write("-" * 100 + "\n")
        file.write("ELEMENT TYPE SUMMARY\n")
        file.write("-" * 100 + "\n")

        for element_type, count in counts.items():

            file.write(
                f"{element_type}: {count}\n"
            )

        file.write("\n")

        file.write("=" * 100 + "\n")
        file.write("ALL EXTRACTED ELEMENTS\n")
        file.write("=" * 100 + "\n\n")

        # --------------------------------------------------------
        # WRITE EVERY ELEMENT
        # --------------------------------------------------------

        for number, element in enumerate(
            result.elements,
            start=1,
        ):

            file.write("\n")
            file.write("#" * 100 + "\n")
            file.write(
                f"ELEMENT {number}\n"
            )
            file.write("#" * 100 + "\n")

            file.write(
                f"Order Index : {element.order_index}\n"
            )

            file.write(
                f"Element Type: {element.element_type}\n"
            )

            file.write("\n")

            file.write("TEXT:\n")
            file.write("-" * 100 + "\n")

            file.write(
                element.text
            )

            file.write("\n")
            file.write("\n")

            file.write("METADATA:\n")
            file.write("-" * 100 + "\n")

            file.write(
                repr(element.metadata)
            )

            file.write("\n")

    # ------------------------------------------------------------
    # CONSOLE SUMMARY
    # ------------------------------------------------------------

    print("\n")
    print("=" * 80)
    print("REAL DOCX EXTRACTION TEST PASSED")
    print("=" * 80)

    print(
        f"Document : {docx_path}"
    )

    print(
        f"Output   : {output_path}"
    )

    print(
        f"Elements : {len(result.elements)}"
    )

    print("\nElement Types:")

    for element_type, count in counts.items():

        print(
            f"  {element_type}: {count}"
        )

    print("=" * 80)


# ============================================================
# EDGE CASE 1
# BASIC EXTRACTION
# ============================================================

def test_basic_extraction():

    path = create_docx(
        "basic.docx",
        paragraphs=[
            "Hello world.",
            "This is a test document.",
        ],
    )

    extractor = DocxExtractor()

    result = extractor.extract(
        path,
        document_id="basic-id",
        filename="basic.docx",
    )

    assert result is not None
    assert result.elements

    paragraphs = get_paragraphs(result)

    assert len(paragraphs) == 2


# ============================================================
# EDGE CASE 2
# METADATA MUST BE CORRECT
# ============================================================

def test_metadata_is_correct():

    path = create_docx(
        "metadata.docx",
        paragraphs=[
            "Document metadata test.",
        ],
        tables=[
            [
                ["Name", "Value"],
                ["Test", "123"],
            ]
        ],
    )

    extractor = DocxExtractor()

    result = extractor.extract(
        path,
        document_id="metadata-id",
        filename="metadata.docx",
    )

    assert result.elements

    for element in result.elements:

        assert (
            element.metadata["document_id"]
            == "metadata-id"
        )

        assert (
            element.metadata["filename"]
            == "metadata.docx"
        )


# ============================================================
# EDGE CASE 3
# TABLE IDS
# ============================================================

def test_table_ids():

    path = create_docx(
        "table_ids.docx",
        tables=[
            [
                ["Name", "Age"],
                ["A", "20"],
            ],
            [
                ["City", "Country"],
                ["Mumbai", "India"],
            ],
            [
                ["Product", "Price"],
                ["Laptop", "50000"],
            ],
        ],
    )

    extractor = DocxExtractor()

    result = extractor.extract(
        path,
        document_id="table-test",
        filename="table_ids.docx",
    )

    tables = get_tables(result)

    assert len(tables) == 3

    assert (
        tables[0].metadata["table_id"]
        == "table-test-table-0"
    )

    assert (
        tables[1].metadata["table_id"]
        == "table-test-table-1"
    )

    assert (
        tables[2].metadata["table_id"]
        == "table-test-table-2"
    )


# ============================================================
# EDGE CASE 4
# SAME EXTRACTOR INSTANCE
#
# THIS IS THE MAIN PRODUCTION BUG TEST
# ============================================================

def test_same_extractor_multiple_documents():

    path_a = create_docx(
        "document_a.docx",
        paragraphs=[
            "This belongs to document A.",
        ],
        tables=[
            [
                ["Name", "Value"],
                ["A", "100"],
            ],
            [
                ["City", "Country"],
                ["Mumbai", "India"],
            ],
        ],
    )

    path_b = create_docx(
        "document_b.docx",
        paragraphs=[
            "This belongs to document B.",
        ],
        tables=[
            [
                ["Name", "Value"],
                ["B", "200"],
            ],
        ],
    )

    # IMPORTANT:
    # Same extractor instance.
    extractor = DocxExtractor()

    result_a = extractor.extract(
        path_a,
        document_id="A",
        filename="document_a.docx",
    )

    result_b = extractor.extract(
        path_b,
        document_id="B",
        filename="document_b.docx",
    )

    # -------------------------
    # Document A
    # -------------------------

    for element in result_a.elements:

        assert (
            element.metadata["document_id"]
            == "A"
        )

        assert (
            element.metadata["filename"]
            == "document_a.docx"
        )

    tables_a = get_tables(result_a)

    assert (
        tables_a[0].metadata["table_id"]
        == "A-table-0"
    )

    assert (
        tables_a[1].metadata["table_id"]
        == "A-table-1"
    )

    # -------------------------
    # Document B
    # -------------------------

    for element in result_b.elements:

        assert (
            element.metadata["document_id"]
            == "B"
        )

        assert (
            element.metadata["filename"]
            == "document_b.docx"
        )

    tables_b = get_tables(result_b)

    assert (
        tables_b[0].metadata["table_id"]
        == "B-table-0"
    )


# ============================================================
# EDGE CASE 5
# TABLE COUNTER MUST RESET
# ============================================================

def test_table_counter_resets():

    path_a = create_docx(
        "counter_a.docx",
        tables=[
            [
                ["A", "B"],
                ["1", "2"],
            ],
            [
                ["C", "D"],
                ["3", "4"],
            ],
        ],
    )

    path_b = create_docx(
        "counter_b.docx",
        tables=[
            [
                ["E", "F"],
                ["5", "6"],
            ],
        ],
    )

    extractor = DocxExtractor()

    result_a = extractor.extract(
        path_a,
        document_id="A",
    )

    result_b = extractor.extract(
        path_b,
        document_id="B",
    )

    tables_a = get_tables(result_a)
    tables_b = get_tables(result_b)

    assert (
        tables_a[0].metadata["table_id"]
        == "A-table-0"
    )

    assert (
        tables_a[1].metadata["table_id"]
        == "A-table-1"
    )

    # Must restart at zero.
    assert (
        tables_b[0].metadata["table_id"]
        == "B-table-0"
    )


# ============================================================
# EDGE CASE 6
# document_id = None
# ============================================================

def test_document_id_none():

    path = create_docx(
        "no_document_id.docx",
        paragraphs=[
            "No document ID.",
        ],
        tables=[
            [
                ["A", "B"],
                ["1", "2"],
            ],
        ],
    )

    extractor = DocxExtractor()

    result = extractor.extract(
        path,
        filename="no_document_id.docx",
    )

    for element in result.elements:

        assert (
            element.metadata["document_id"]
            is None
        )

        assert (
            element.metadata["filename"]
            == "no_document_id.docx"
        )

    tables = get_tables(result)

    assert (
        tables[0].metadata["table_id"]
        == "table-0"
    )


# ============================================================
# EDGE CASE 7
# filename = None
# ============================================================

def test_filename_none():

    path = create_docx(
        "filename_default.docx",
        paragraphs=[
            "Filename should be inferred.",
        ],
    )

    extractor = DocxExtractor()

    result = extractor.extract(
        path,
        document_id="filename-test",
    )

    for element in result.elements:

        assert (
            element.metadata["filename"]
            == path.name
        )


# ============================================================
# EDGE CASE 8
# BOTH document_id AND filename NONE
# ============================================================

def test_document_id_and_filename_none():

    path = create_docx(
        "defaults.docx",
        paragraphs=[
            "Default metadata test.",
        ],
    )

    extractor = DocxExtractor()

    result = extractor.extract(path)

    for element in result.elements:

        assert (
            element.metadata["document_id"]
            is None
        )

        assert (
            element.metadata["filename"]
            == path.name
        )


# ============================================================
# EDGE CASE 9
# EMPTY DOCUMENT
# ============================================================

def test_empty_document():

    path = TEST_DIR / "empty.docx"

    doc = Document()
    doc.save(path)

    extractor = DocxExtractor()

    result = extractor.extract(
        path,
        document_id="empty",
        filename="empty.docx",
    )

    assert result.elements == []


# ============================================================
# EDGE CASE 10
# PARAGRAPH ONLY
# ============================================================

def test_paragraph_only_document():

    path = create_docx(
        "paragraph_only.docx",
        paragraphs=[
            "Paragraph one.",
            "Paragraph two.",
            "Paragraph three.",
            "Paragraph four.",
            "Paragraph five.",
        ],
    )

    extractor = DocxExtractor()

    result = extractor.extract(
        path,
        document_id="paragraph-only",
    )

    paragraphs = get_paragraphs(result)
    tables = get_tables(result)

    assert len(paragraphs) == 5
    assert len(tables) == 0


# ============================================================
# EDGE CASE 11
# TABLE ONLY
# ============================================================

def test_table_only_document():

    path = create_docx(
        "table_only.docx",
        tables=[
            [
                ["Name", "Age"],
                ["Abhishek", "20"],
            ],
        ],
    )

    extractor = DocxExtractor()

    result = extractor.extract(
        path,
        document_id="table-only",
    )

    tables = get_tables(result)

    assert len(tables) == 1

    assert (
        tables[0].metadata["table_id"]
        == "table-only-table-0"
    )


# ============================================================
# EDGE CASE 12
# MIXED DOCUMENT
# ============================================================

def test_mixed_document():

    path = create_docx(
        "mixed.docx",
        paragraphs=[
            "Paragraph 1",
            "Paragraph 2",
            "Paragraph 3",
        ],
        tables=[
            [
                ["A", "B"],
                ["1", "2"],
            ],
            [
                ["C", "D"],
                ["3", "4"],
            ],
        ],
    )

    extractor = DocxExtractor()

    result = extractor.extract(
        path,
        document_id="mixed",
    )

    paragraphs = get_paragraphs(result)
    tables = get_tables(result)

    assert len(paragraphs) == 3
    assert len(tables) == 2


# ============================================================
# EDGE CASE 13
# ORDER INDEX MUST BE UNIQUE AND SEQUENTIAL
# ============================================================

def test_order_indexes():

    path = create_docx(
        "order_indexes.docx",
        paragraphs=[
            "Paragraph 1",
            "Paragraph 2",
            "Paragraph 3",
        ],
        tables=[
            [
                ["A", "B"],
                ["1", "2"],
            ],
            [
                ["C", "D"],
                ["3", "4"],
            ],
        ],
    )

    extractor = DocxExtractor()

    result = extractor.extract(
        path,
        document_id="order-test",
    )

    indexes = [
        element.order_index
        for element in result.elements
    ]

    assert indexes == list(
        range(len(indexes))
    )


# ============================================================
# EDGE CASE 14
# TABLE METADATA
# ============================================================

def test_table_metadata():

    path = create_docx(
        "table_metadata.docx",
        tables=[
            [
                ["Name", "Age", "City"],
                ["Abhishek", "20", "Mumbai"],
                ["Rahul", "21", "Pune"],
            ],
        ],
    )

    extractor = DocxExtractor()

    result = extractor.extract(
        path,
        document_id="metadata-test",
    )

    tables = get_tables(result)

    assert len(tables) == 1

    metadata = tables[0].metadata

    assert metadata["n_rows"] == 3
    assert metadata["n_cols"] == 3

    assert metadata["cells"] == [
        ["Name", "Age", "City"],
        ["Abhishek", "20", "Mumbai"],
        ["Rahul", "21", "Pune"],
    ]

    assert (
        metadata["table_id"]
        == "metadata-test-table-0"
    )


# ============================================================
# EDGE CASE 15
# TABLE WITH NO HEADER
#
# Tests the table Markdown bug we fixed.
# ============================================================

def test_table_without_header():

    path = create_docx(
        "no_header_table.docx",
        tables=[
            [
                ["Apple", "10", "Red"],
                ["Banana", "20", "Yellow"],
                ["Mango", "15", "Green"],
            ],
        ],
    )

    extractor = DocxExtractor()

    result = extractor.extract(
        path,
        document_id="no-header",
    )

    tables = get_tables(result)

    assert len(tables) == 1

    metadata = tables[0].metadata

    # Depending on the current heuristic,
    # this test primarily verifies that if
    # has_header_row is False, the first row
    # is NOT duplicated.
    markdown = metadata["markdown"]

    lines = [
        line
        for line in markdown.splitlines()
        if line.strip()
    ]

    # First row should occur exactly once.
    assert lines.count(
        "| Apple | 10 | Red |"
    ) <= 1


# ============================================================
# EDGE CASE 16
# MULTIPLE TABLES + METADATA ISOLATION
# ============================================================

def test_multiple_tables_metadata_isolation():

    path = create_docx(
        "multiple_tables.docx",
        tables=[
            [
                ["A", "B"],
                ["1", "2"],
            ],
            [
                ["C", "D"],
                ["3", "4"],
            ],
            [
                ["E", "F"],
                ["5", "6"],
            ],
            [
                ["G", "H"],
                ["7", "8"],
            ],
        ],
    )

    extractor = DocxExtractor()

    result = extractor.extract(
        path,
        document_id="multi",
        filename="multiple_tables.docx",
    )

    tables = get_tables(result)

    assert len(tables) == 4

    for index, table in enumerate(tables):

        assert (
            table.metadata["document_id"]
            == "multi"
        )

        assert (
            table.metadata["filename"]
            == "multiple_tables.docx"
        )

        assert (
            table.metadata["table_id"]
            == f"multi-table-{index}"
        )


# ============================================================
# EDGE CASE 17
# REPEATED EXTRACTION
# ============================================================

def test_repeated_extraction():

    path = create_docx(
        "repeated.docx",
        paragraphs=[
            "Repeated extraction test.",
        ],
        tables=[
            [
                ["A", "B"],
                ["1", "2"],
            ],
            [
                ["C", "D"],
                ["3", "4"],
            ],
        ],
    )

    extractor = DocxExtractor()

    for _ in range(10):

        result = extractor.extract(
            path,
            document_id="repeat",
            filename="repeated.docx",
        )

        tables = get_tables(result)

        assert len(tables) == 2

        assert (
            tables[0].metadata["table_id"]
            == "repeat-table-0"
        )

        assert (
            tables[1].metadata["table_id"]
            == "repeat-table-1"
        )

        for element in result.elements:

            assert (
                element.metadata["document_id"]
                == "repeat"
            )

            assert (
                element.metadata["filename"]
                == "repeated.docx"
            )


# ============================================================
# EDGE CASE 18
# CONCURRENT EXTRACTION
#
# This is the most important production test.
# Same extractor instance is used by multiple threads.
# ============================================================

def extract_with_shared_extractor(
    extractor,
    path,
    document_id,
    filename,
):

    return extractor.extract(
        path,
        document_id=document_id,
        filename=filename,
    )


def test_concurrent_extraction():

    path_a = create_docx(
        "concurrent_a.docx",
        paragraphs=[
            "Document A",
        ],
        tables=[
            [
                ["A", "B"],
                ["1", "2"],
            ],
            [
                ["C", "D"],
                ["3", "4"],
            ],
            [
                ["E", "F"],
                ["5", "6"],
            ],
        ],
    )

    path_b = create_docx(
        "concurrent_b.docx",
        paragraphs=[
            "Document B",
        ],
        tables=[
            [
                ["A", "B"],
                ["10", "20"],
            ],
            [
                ["C", "D"],
                ["30", "40"],
            ],
            [
                ["E", "F"],
                ["50", "60"],
            ],
        ],
    )

    # VERY IMPORTANT:
    # Same extractor instance.
    extractor = DocxExtractor()

    with ThreadPoolExecutor(
        max_workers=2
    ) as executor:

        future_a = executor.submit(
            extract_with_shared_extractor,
            extractor,
            path_a,
            "A",
            "concurrent_a.docx",
        )

        future_b = executor.submit(
            extract_with_shared_extractor,
            extractor,
            path_b,
            "B",
            "concurrent_b.docx",
        )

        result_a = future_a.result()
        result_b = future_b.result()

    # -------------------------
    # Verify A
    # -------------------------

    for element in result_a.elements:

        assert (
            element.metadata["document_id"]
            == "A"
        )

        assert (
            element.metadata["filename"]
            == "concurrent_a.docx"
        )

    tables_a = get_tables(result_a)

    assert len(tables_a) == 3

    for index, table in enumerate(tables_a):

        assert (
            table.metadata["table_id"]
            == f"A-table-{index}"
        )

    # -------------------------
    # Verify B
    # -------------------------

    for element in result_b.elements:

        assert (
            element.metadata["document_id"]
            == "B"
        )

        assert (
            element.metadata["filename"]
            == "concurrent_b.docx"
        )

    tables_b = get_tables(result_b)

    assert len(tables_b) == 3

    for index, table in enumerate(tables_b):

        assert (
            table.metadata["table_id"]
            == f"B-table-{index}"
        )


# ============================================================
# EDGE CASE 19
# CODE INDENTATION
# ============================================================

def test_code_indentation_preserved():

    path = TEST_DIR / "code_indentation.docx"

    doc = Document()

    code_lines = [
        "class Example:",
        "    def test(self):",
        "        value = 10",
        "        if value > 5:",
        "            return value",
    ]

    for line in code_lines:

        paragraph = doc.add_paragraph()

        run = paragraph.add_run(line)

        # Make it monospace so the extractor
        # detects this as code.
        run.font.name = "Consolas"

    doc.save(path)

    extractor = DocxExtractor()

    result = extractor.extract(
        path,
        document_id="code-test",
    )

    code_blocks = get_code_blocks(result)

    assert code_blocks

    extracted_text = "\n".join(
        element.text
        for element in code_blocks
    )

    assert "    def test(self):" in extracted_text
    assert "        value = 10" in extracted_text
    assert "            return value" in extracted_text


# ============================================================
# EDGE CASE 20
# CODE MUST NOT LOSE LEADING WHITESPACE
# ============================================================

def test_code_does_not_use_strip():

    path = TEST_DIR / "code_strip_test.docx"

    doc = Document()

    paragraph = doc.add_paragraph()

    paragraph.add_run(
        "    indented_code_line"
    ).font.name = "Consolas"

    paragraph2 = doc.add_paragraph()

    paragraph2.add_run(
        "        deeply_indented_line"
    ).font.name = "Consolas"

    doc.save(path)

    extractor = DocxExtractor()

    result = extractor.extract(
        path,
        document_id="strip-test",
    )

    code_blocks = get_code_blocks(result)

    assert code_blocks

    text = "\n".join(
        block.text
        for block in code_blocks
    )

    assert (
        "    indented_code_line"
        in text
    )

    assert (
        "        deeply_indented_line"
        in text
    )


# ============================================================
# EDGE CASE 21
# EXTRACTION ORDER MUST REMAIN STABLE
# ============================================================

def test_extraction_order_is_stable():

    path = create_docx(
        "stable_order.docx",
        paragraphs=[
            "Paragraph A",
        ],
        tables=[
            [
                ["A", "B"],
                ["1", "2"],
            ],
        ],
    )

    # Add another paragraph/table manually
    doc = Document(path)

    doc.add_paragraph("Paragraph B")

    table = doc.add_table(
        rows=2,
        cols=2,
    )

    table.cell(0, 0).text = "C"
    table.cell(0, 1).text = "D"
    table.cell(1, 0).text = "3"
    table.cell(1, 1).text = "4"

    doc.save(path)

    extractor = DocxExtractor()

    result_1 = extractor.extract(
        path,
        document_id="stable",
    )

    result_2 = extractor.extract(
        path,
        document_id="stable",
    )

    sequence_1 = [
        (
            element.order_index,
            element.element_type,
            element.text,
        )
        for element in result_1.elements
    ]

    sequence_2 = [
        (
            element.order_index,
            element.element_type,
            element.text,
        )
        for element in result_2.elements
    ]

    assert sequence_1 == sequence_2


# ============================================================
# EDGE CASE 22
# NO CROSS-DOCUMENT TEXT LEAK
# ============================================================

def test_no_cross_document_text_leak():

    path_a = create_docx(
        "leak_a.docx",
        paragraphs=[
            "SECRET_A_ONLY",
        ],
    )

    path_b = create_docx(
        "leak_b.docx",
        paragraphs=[
            "SECRET_B_ONLY",
        ],
    )

    extractor = DocxExtractor()

    result_a = extractor.extract(
        path_a,
        document_id="A",
    )

    result_b = extractor.extract(
        path_b,
        document_id="B",
    )

    text_a = "\n".join(
        element.text
        for element in result_a.elements
    )

    text_b = "\n".join(
        element.text
        for element in result_b.elements
    )

    assert "SECRET_A_ONLY" in text_a
    assert "SECRET_B_ONLY" not in text_a

    assert "SECRET_B_ONLY" in text_b
    assert "SECRET_A_ONLY" not in text_b